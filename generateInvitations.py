#! python3
# generateInvitations.py

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

guestFile = open('guests.txt')
guests = guestFile.readlines()
guestFile.close()
doc = docx.Document()
for guest in guests:
    paraObj1 = doc.add_paragraph('It would be a pleasure to have the company of')
    paraObj1.runs[0].italic = True
    paraObj1.runs[0].bold = True
    paraObj1.runs[0].font.name = 'Droid Serif'
    paraObj1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paraObj2 = doc.add_paragraph(guest.rstrip())
    paraObj2.runs[0].font.name = 'Latin Modern Roman'
    paraObj2.runs[0].font.size = docx.shared.Pt(20)
    paraObj2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paraObj3 = doc.add_paragraph('at')
    paraObj3.runs[0].underline = True
    paraObj3.runs[0].italic = True
    paraObj3.runs[0].bold = True
    paraObj3.runs[0].font.name = 'Droid Serif'
    paraObj3.add_run(' 11011 Memory Lane on the Evening of')
    paraObj3.runs[1].italic = True
    paraObj3.runs[1].bold = True    
    paraObj3.runs[1].font.name = 'Droid Serif'
    paraObj3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paraObj4 = doc.add_paragraph('April 1st')
    paraObj4.runs[0].font.name = 'Source Sans Pro Light'
    paraObj4.runs[0].font.size = docx.shared.Pt(14)
    paraObj4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paraObj5 = doc.add_paragraph('at')
    paraObj5.runs[0].underline = True
    paraObj5.runs[0].italic = True
    paraObj5.runs[0].bold = True
    paraObj5.runs[0].font.name = 'Droid Serif'
    paraObj5.add_run(' 7 o\' clock')
    paraObj5.runs[1].italic = True
    paraObj5.runs[1].bold = True
    paraObj5.runs[1].font.name = 'Droid Serif'
    paraObj5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
doc.save('guestInvitations.doc')

    
